import streamlit as st
import tempfile
import base64
import requests
import PyPDF2
import io
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from streamlit_option_menu import option_menu
import streamlit as st
import os
from dotenv import load_dotenv

# Configuration de la page
st.set_page_config(
    page_title="Assistant IA - Générateur des Cas de test",
    page_icon="📋",
    layout="wide"
)

# Essayer de charger depuis .env pour le développement local
load_dotenv()

# Récupérer la clé API (priorité aux secrets Streamlit)
def get_api_key():
    # D'abord chercher dans les secrets Streamlit (production)
    if 'ANTHROPIC_API_KEY' in st.secrets:
        return st.secrets['ANTHROPIC_API_KEY']
    # Sinon chercher dans les variables d'environnement (dev local)
    else:
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            st.error("Clé API Anthropic non trouvée. Vérifiez votre configuration.")
        return api_key
# Fonction pour extraire le texte d'un fichier DOCX
def extract_text_from_docx(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
    text = []
    for page in pdf_reader.pages:
        text.append(page.extract_text())
    return "\n".join(text)

# Fonction pour créer un lien de téléchargement
def create_download_link(file_data, file_type, file_name):
    b64 = base64.b64encode(file_data).decode()
    
    if file_type == "pdf":
        href = f'<a href="data:application/pdf;base64,{b64}" download="{file_name}" class="css-1offfwp edgvbvh9" style="background-color: #0066cc; color: white; font-weight: bold; padding: 0.5rem; border-radius: 0.25rem; text-decoration: none; width: 100%; display: inline-block; text-align: center;">Télécharger en PDF</a>'
    else:
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{file_name}" class="css-1offfwp edgvbvh9" style="background-color: #0066cc; color: white; font-weight: bold; padding: 0.5rem; border-radius: 0.25rem; text-decoration: none; width: 100%; display: inline-block; text-align: center;">Télécharger en DOCX</a>'
    
    return href
# Fonction pour créer un PDF bien structuré à partir du texte
def create_pdf(output_text, filename="cas_de_test.pdf"):
    buffer = io.BytesIO()
    
    # Définition des marges et de la taille
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=60,
        leftMargin=60,
        topMargin=72,
        bottomMargin=60
    )
    
    # Styles améliorés
    styles = getSampleStyleSheet()
    
    # Style du titre principal
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=22,
        textColor=colors.darkblue,
        spaceAfter=20,
        alignment=TA_CENTER,
        borderWidth=0,
        borderPadding=10,
        leading=30
    )
    
    # Style pour les titres de sections
    heading2_style = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.darkblue,
        spaceAfter=12,
        spaceBefore=12,
        borderBottom=1,
        borderColor=colors.darkblue,
        paddingBottom=5,
        leading=20
    )
    
    # Style pour les scénarios
    heading3_style = ParagraphStyle(
        'Heading3',
        parent=styles['Heading3'],
        fontSize=14,
        textColor=colors.navy,
        spaceAfter=8,
        spaceBefore=8,
        leading=16,
        bulletIndent=0
    )
    
    # Style pour le texte normal
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        leading=16,  # Espacement entre les lignes
        firstLineIndent=0
    )
    
    # Style pour les préconditions
    precondition_style = ParagraphStyle(
        'Precondition',
        parent=normal_style,
        leftIndent=20,
        textColor=colors.black
    )
    
    # Style pour les étapes
    steps_style = ParagraphStyle(
        'Steps',
        parent=normal_style,
        leftIndent=20,
        textColor=colors.black
    )
    
    # Style pour les résultats attendus
    result_style = ParagraphStyle(
        'Result',
        parent=normal_style,
        leftIndent=20,
        textColor=colors.black,
        backColor=colors.lightgrey,
        borderPadding=5
    )
    
    # Style pour la date
    date_style = ParagraphStyle(
        'Date',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.grey,
        alignment=TA_CENTER
    )
    
    flowables = []
    
    # Page de titre
    flowables.append(Paragraph("Cas de Test", title_style))
    flowables.append(Spacer(1, 0.3*inch))
    
    # Date de génération
    from datetime import datetime
    date_str = datetime.now().strftime("Généré le %d/%m/%Y à %H:%M")
    flowables.append(Paragraph(date_str, date_style))
    flowables.append(Spacer(1, 0.5*inch))
    
    # Splitter le texte en sections et paragraphes
    lines = output_text.split('\n')
    current_section = []
    in_precondition = False
    in_steps = False
    in_result = False
    
    for line in lines:
        # Gestion des sections principales (exigences)
        if line.startswith('## ') or line.startswith('**EXIGENCE ') or (line.startswith('EXIGENCE ') and not line.startswith('EXIGENCE:')):
            # Si on a une section précédente, on la traite
            if current_section:
                section_text = '\n'.join(current_section)
                para = Paragraph(section_text.replace('\n', '<br/>'), normal_style)
                flowables.append(para)
                flowables.append(Spacer(1, 0.15*inch))
                current_section = []
            
            # Nouvelle section avec saut de page
            flowables.append(PageBreak())
            
            # Titre de la section (exigence)
            if line.startswith('## '):
                flowables.append(Paragraph(line[3:], heading2_style))
            else:
                if line.startswith('**EXIGENCE '):
                    clean_line = line.replace('**', '')
                else:
                    clean_line = line
                flowables.append(Paragraph(clean_line, heading2_style))
                
        # Gestion des scénarios
        elif line.startswith('**Scenario ') or line.startswith('Scenario ') or line.startswith('**Cas '):
            # Si on a une section précédente, on la traite
            if current_section:
                section_text = '\n'.join(current_section)
                
                # Appliquer le style approprié
                if in_precondition:
                    style_to_use = precondition_style
                    in_precondition = False
                elif in_steps:
                    style_to_use = steps_style
                    in_steps = False
                elif in_result:
                    style_to_use = result_style
                    in_result = False
                else:
                    style_to_use = normal_style
                
                para = Paragraph(section_text.replace('\n', '<br/>'), style_to_use)
                flowables.append(para)
                flowables.append(Spacer(1, 0.15*inch))
                current_section = []
            
            # Nouveau scénario (sous-section)
            clean_line = line.replace('**', '')
            flowables.append(Paragraph(clean_line, heading3_style))
            flowables.append(Spacer(1, 0.1*inch))
            
        # Gestion des préconditions, étapes et résultats
        elif "Précondition" in line or "Precondition" in line or line.startswith("**Précondition") or line.startswith("**Precondition"):
            if current_section:
                section_text = '\n'.join(current_section)
                para = Paragraph(section_text.replace('\n', '<br/>'), normal_style)
                flowables.append(para)
                current_section = []
            
            # Nouvelle précondition
            clean_line = line.replace('**', '')
            current_section.append(clean_line)
            in_precondition = True
            in_steps = False
            in_result = False
            
        elif "Etapes" in line or "Étapes" in line or line.startswith("**Etapes") or line.startswith("**Étapes"):
            if current_section:
                section_text = '\n'.join(current_section)
                para = Paragraph(section_text.replace('\n', '<br/>'), precondition_style if in_precondition else normal_style)
                flowables.append(para)
                current_section = []
            
            # Nouvelles étapes
            clean_line = line.replace('**', '')
            current_section.append(clean_line)
            in_precondition = False
            in_steps = True
            in_result = False
            
        elif "Résultat attendu" in line or "Résultats attendus" in line or line.startswith("**Résultat"):
            if current_section:
                section_text = '\n'.join(current_section)
                para = Paragraph(section_text.replace('\n', '<br/>'), steps_style if in_steps else normal_style)
                flowables.append(para)
                current_section = []
            
            # Nouveau résultat attendu
            clean_line = line.replace('**', '')
            current_section.append(clean_line)
            in_precondition = False
            in_steps = False
            in_result = True
            
        else:
            current_section.append(line)
    
    # Traiter la dernière section
    if current_section:
        section_text = '\n'.join(current_section)
        
        # Appliquer le style approprié pour la dernière section
        if in_precondition:
            style_to_use = precondition_style
        elif in_steps:
            style_to_use = steps_style
        elif in_result:
            style_to_use = result_style
        else:
            style_to_use = normal_style
        
        para = Paragraph(section_text.replace('\n', '<br/>'), style_to_use)
        flowables.append(para)
    
    # Génération du PDF
    doc.build(flowables)
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data
# Fonction pour créer un DOCX à partir du texte
def create_docx(output_text, filename="cas_de_test.docx"):
    doc = docx.Document()
    
    # Définition des styles
    styles = doc.styles
    
    # Style pour le titre principal
    title_style = doc.styles.add_style('Title Custom', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = docx.shared.Pt(22)
    title_style.font.color.rgb = docx.shared.RGBColor(0, 51, 102)  # Dark blue
    title_style.font.bold = True
    title_style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = docx.shared.Pt(20)
    
    # Style pour les titres de section (exigences)
    heading1_style = doc.styles.add_style('Heading1 Custom', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = docx.shared.Pt(16)
    heading1_style.font.color.rgb = docx.shared.RGBColor(0, 51, 102)  # Dark blue
    heading1_style.font.bold = True
    heading1_style.paragraph_format.space_before = docx.shared.Pt(12)
    heading1_style.paragraph_format.space_after = docx.shared.Pt(8)
    
    # Style pour les scénarios
    heading2_style = doc.styles.add_style('Heading2 Custom', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = docx.shared.Pt(14)
    heading2_style.font.color.rgb = docx.shared.RGBColor(0, 51, 102)  # Dark blue
    heading2_style.font.bold = True
    heading2_style.paragraph_format.space_before = docx.shared.Pt(10)
    heading2_style.paragraph_format.space_after = docx.shared.Pt(6)
    
    # Style pour les préconditions, étapes, résultats attendus
    section_title_style = doc.styles.add_style('Section Title', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    section_title_style.font.size = docx.shared.Pt(12)
    section_title_style.font.bold = True
    section_title_style.paragraph_format.space_before = docx.shared.Pt(6)
    section_title_style.paragraph_format.space_after = docx.shared.Pt(2)
    
    # Style pour le contenu normal
    normal_style = doc.styles.add_style('Normal Custom', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = docx.shared.Pt(11)
    normal_style.paragraph_format.space_after = docx.shared.Pt(6)
    normal_style.paragraph_format.left_indent = docx.shared.Pt(20)
    
    # Style pour les résultats attendus (avec fond gris)
    result_style = doc.styles.add_style('Result Custom', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    result_style.font.size = docx.shared.Pt(11)
    result_style.paragraph_format.space_after = docx.shared.Pt(6)
    result_style.paragraph_format.left_indent = docx.shared.Pt(20)
    
    # Ajouter un titre
    title = doc.add_paragraph("Cas de Test", style='Title Custom')
    
    # Ajouter la date
    from datetime import datetime
    date_paragraph = doc.add_paragraph()
    date_paragraph.add_run(datetime.now().strftime("Généré le %d/%m/%Y à %H:%M"))
    date_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.runs[0].font.size = docx.shared.Pt(10)
    date_paragraph.runs[0].font.color.rgb = docx.shared.RGBColor(128, 128, 128)  # Gris
    
    # Ajouter un saut de page après le titre
    doc.add_page_break()
    
    # Traiter les différentes sections
    lines = output_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Exigence ou titre de section principal
        if line.startswith('## ') or line.startswith('**EXIGENCE ') or (line.startswith('EXIGENCE ') and not line.startswith('EXIGENCE:')):
            # Nettoyage du texte
            if line.startswith('## '):
                clean_title = line[3:]
            elif line.startswith('**EXIGENCE '):
                clean_title = line.replace('**', '')
            else:
                clean_title = line
            
            doc.add_paragraph(clean_title, style='Heading1 Custom')
            
        # Scénario
        elif line.startswith('**Scenario ') or line.startswith('Scenario ') or line.startswith('**Cas ') or 'Scenario' in line:
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line, style='Heading2 Custom')
            
        # Précondition
        elif "Précondition" in line or "Precondition" in line:
            # Séparer le titre de la précondition et son contenu
            if ":" in line:
                parts = line.split(":", 1)
                title_part = parts[0].replace('**', '')
                content_part = parts[1].strip() if len(parts) > 1 else ""
                
                # Ajouter le titre de la précondition
                doc.add_paragraph(title_part + ":", style='Section Title')
                
                # Ajouter le contenu de la précondition si présent
                if content_part:
                    doc.add_paragraph(content_part, style='Normal Custom')
                
                # Chercher les lignes suivantes qui font partie de la précondition
                i += 1
                while i < len(lines) and not (lines[i].strip().startswith("**") or "Etapes" in lines[i] or "Étapes" in lines[i] or "Résultat" in lines[i]):
                    if lines[i].strip():
                        doc.add_paragraph(lines[i].strip(), style='Normal Custom')
                    i += 1
                i -= 1  # Revenir d'une ligne pour que le prochain titre soit traité
            else:
                # Si pas de ":", traiter comme un titre simple
                doc.add_paragraph(line.replace('**', ''), style='Section Title')
            
        # Étapes
        elif "Etapes" in line or "Étapes" in line:
            # Séparer le titre des étapes et son contenu
            if ":" in line:
                parts = line.split(":", 1)
                title_part = parts[0].replace('**', '')
                content_part = parts[1].strip() if len(parts) > 1 else ""
                
                # Ajouter le titre des étapes
                doc.add_paragraph(title_part + ":", style='Section Title')
                
                # Ajouter le contenu des étapes si présent
                if content_part:
                    doc.add_paragraph(content_part, style='Normal Custom')
                
                # Chercher les lignes suivantes qui font partie des étapes
                i += 1
                while i < len(lines) and not (lines[i].strip().startswith("**") or "Résultat" in lines[i]):
                    if lines[i].strip():
                        doc.add_paragraph(lines[i].strip(), style='Normal Custom')
                    i += 1
                i -= 1  # Revenir d'une ligne pour que le prochain titre soit traité
            else:
                # Si pas de ":", traiter comme un titre simple
                doc.add_paragraph(line.replace('**', ''), style='Section Title')
            
        # Résultat attendu
        elif "Résultat attendu" in line or "Résultats attendus" in line:
            # Séparer le titre du résultat et son contenu
            if ":" in line:
                parts = line.split(":", 1)
                title_part = parts[0].replace('**', '')
                content_part = parts[1].strip() if len(parts) > 1 else ""
                
                # Ajouter le titre du résultat
                doc.add_paragraph(title_part + ":", style='Section Title')
                
                # Ajouter le contenu du résultat si présent
                if content_part:
                    result_para = doc.add_paragraph(content_part, style='Result Custom')
                    # Ajouter un fond gris (simulation)
                    for run in result_para.runs:
                        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GRAY_25
                
                # Chercher les lignes suivantes qui font partie du résultat
                i += 1
                while i < len(lines) and not (lines[i].strip().startswith("**") or "Scenario" in lines[i] or "Cas" in lines[i] or "EXIGENCE" in lines[i] or lines[i].strip().startswith("##")):
                    if lines[i].strip():
                        result_para = doc.add_paragraph(lines[i].strip(), style='Result Custom')
                        # Ajouter un fond gris (simulation)
                        for run in result_para.runs:
                            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GRAY_25
                    i += 1
                i -= 1  # Revenir d'une ligne pour que le prochain titre soit traité
            else:
                # Si pas de ":", traiter comme un titre simple
                doc.add_paragraph(line.replace('**', ''), style='Section Title')
            
        # Ligne normale
        elif line:
            doc.add_paragraph(line, style='Normal')
        
        i += 1
    
    # Sauvegarder dans un BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_data = buffer.getvalue()
    buffer.close()
    
    return docx_data
# Fonction pour générer des cas de test via l'API Claude
def generate_test_cases(requirements, format_type, context="", example_case=""):
    import os
    from dotenv import load_dotenv
    import langdetect
    
    # Charger les variables d'environnement
    load_dotenv()
    
    # Récupérer la clé API depuis les variables d'environnement
    api_key = get_api_key()
    
    if not api_key:
        return "Erreur: Clé API non trouvée dans les variables d'environnement. Vérifiez votre fichier .env"
    
    # Configuration pour l'API
    headers = {
        "anthropic-api-key": api_key,
        "content-type": "application/json",
        "anthropic-version": "2023-06-01",
        "x-api-key": api_key
    }
    
    # Détection de la langue de l'input (contexte + exigences)
    try:
        input_text = context + " " + requirements
        detected_lang = langdetect.detect(input_text)
        
        # Par défaut, on considère le français
        lang = "fr"
        
        # Si l'anglais est détecté, on adapte les instructions
        if detected_lang == "en":
            lang = "en"
    except:
        # En cas d'erreur de détection, on reste en français par défaut
        lang = "fr"
    
    # Exemples de format selon la langue
    if lang == "fr":
        exemple_format_defaut = """
**Cas fonctionnels
**Scenario (1) : Connexion OK avec des identifiants valides.
Précondition :** L'utilisateur est inscrit avec un e-Mail valide et un MP.
**Etapes :**
    1. Accéder à la page de connexion.
    2. Saisir l'e-Mail et le MP valides.
    3. Cliquer sur "Se connecter".
**Résultat attendu** : L'utilisateur est redirigé vers la page d'accueil.

**Scenario (2) : Connexion KO avec un e-Mail valide et un MP invalide.
Précondition :** L'utilisateur est inscrit avec un e-Mail valide et un MP.
**Etapes :**
1. Accéder à la page de connexion.
2. Saisir un e-Mail valide et un MP invalide.
3. Cliquer sur "Se connecter".
**Résultat attendu** : Un message d'erreur s'affiche.

**Scenario (3) : Connexion KO avec un e-Mail invalide et un MP valide.
Précondition :** L'utilisateur est inscrit avec un e-Mail valide et un MP.**
Etapes :**
1. Accéder à la page de connexion.
2. Saisir un e-Mail invalide et un MP valide.
3. Cliquer sur "Se connecter".
**Résultat attendu** : Un message d'erreur s'affiche.

**Cas non-fonctionnels
**Scenario (4) : Connexion KO après plusieurs tentatives**
**Etapes :**
1. Accéder à la page de connexion.
2. Saisir un e-Mail et un MP erronés.
3. Cliquer sur "Se connecter".
4. Refaire les 3 actions plusieurs fois.
5. Un message d'erreur s'affiche et le compte est temporairement bloqué."""
    else:  # Pour l'anglais
        exemple_format_defaut = """
**Functional Test Cases
**Scenario (1): Successful login with valid credentials.
Precondition:** User is registered with a valid email and password.
**Steps:**
    1. Access the login page.
    2. Enter valid email and password.
    3. Click on "Login".
**Expected Result**: User is redirected to the home page.

**Scenario (2): Failed login with valid email and invalid password.
Precondition:** User is registered with a valid email and password.
**Steps:**
1. Access the login page.
2. Enter valid email and invalid password.
3. Click on "Login".
**Expected Result**: An error message is displayed.

**Scenario (3): Failed login with invalid email and valid password.
Precondition:** User is registered with a valid email and password.**
Steps:**
1. Access the login page.
2. Enter invalid email and valid password.
3. Click on "Login".
**Expected Result**: An error message is displayed.

**Non-functional Test Cases
**Scenario (4): Failed login after multiple attempts**
**Steps:**
1. Access the login page.
2. Enter incorrect email and password.
3. Click on "Login".
4. Repeat the 3 actions several times.
5. An error message is displayed and the account is temporarily blocked."""
    
    # Déterminer le format à utiliser
    if format_type == "custom" and example_case.strip():
        exemple_format = example_case
        format_instruction = "Format personnalisé" if lang == "fr" else "Custom format"
    elif format_type == "gherkin":
        format_instruction = "Format Gherkin (Given, When, Then)"
        exemple_format = example_case if example_case.strip() else "Format Gherkin"
    else:
        format_instruction = "Format par défaut" if lang == "fr" else "Default format"
        exemple_format = exemple_format_defaut
    
    # Construire le prompt selon la langue détectée
    if lang == "fr":
        # Construire le prompt en français
        if format_type == "default":
            instruction = f"""
            Génère des cas de test pour l'exigence suivante en utilisant le format par défaut comme dans l'exemple ci-dessous.
            Chaque cas de test doit inclure des scénarios fonctionnels et non-fonctionnels, avec des préconditions, 
            des étapes numérotées , des résultats attendus et donne le max de scenario qui peut couvrir la totalite de l'exigence avec plus de details et prendre en coöpte le moindre exigence de lexigence global.
            
            {"Contexte fonctionnel: " + context if context else ""}
            
            Voici un exemple du format attendu :
            
            {exemple_format}
            
            Maintenant, génère des cas de test détaillés dans ce même format pour l'exigence suivante :
            
            {requirements}
            """
        elif format_type == "gherkin":
            instruction = f"""
            Génère des cas de test pour l'exigence suivante en utilisant le format Gherkin (Given, When, Then).
            Chaque cas de test doit inclure des scénarios fonctionnels et non-fonctionnels et donne le max de scenario qui peut couvrir la totalite de l'exigence.
            
            {"Contexte fonctionnel: " + context if context else ""}
            
            {"Exemple de format attendu: " + example_case if example_case else ""}
            
            Exigence: {requirements}
            """
        else:  # format_type == "custom"
            instruction = f"""
            Génère des cas de test pour l'exigence suivante en utilisant exactement le format personnalisé fourni en exemple.
            Respecte strictement la structure et le style de l'exemple fourni et donne le max de scenario qui peut couvrir la totalite de l'exigence.
            
            {"Contexte fonctionnel: " + context if context else ""}
            
            Voici un exemple du format attendu :
            
            {exemple_format}
            
            Maintenant, génère des cas de test détaillés dans ce même format personnalisé pour l'exigence suivante :
            
            {requirements}
            """
        
        system_prompt = """Tu es un expert en tests logiciels qui génère des cas de test de haute qualité.
        Ton travail consiste à analyser des exigences et à produire des scénarios de test complets, exhaustifs et précis.
        Pour chaque exigence, couvre tous les aspects fonctionnels et non-fonctionnels (performance, sécurité, accessibilité, etc.).

        RÈGLES DE FORMATAGE STRICTES:
        1. Pour chaque SCÉNARIO, commence par indiquer clairement s'il s'agit d'un cas "FONCTIONNEL" ou "NON-FONCTIONNEL".
        2. Utilise des retours à la ligne appropriés pour une meilleure lisibilité.
        3. Les ÉTAPES doivent être numérotées et indentées avec une tabulation, une étape par ligne.
        4. Les PRÉCONDITIONS doivent être clairement séparées du reste.
        5. Les RÉSULTATS ATTENDUS doivent être clairement mis en évidence, en utilisant un retour à la ligne après "Résultat attendu :".
        6. Assure-toi de bien mettre en gras (**) les titres des sections (Scenario, Précondition, Étapes, Résultat attendu).
        7. Génère le MAXIMUM de scénarios possibles pour couvrir TOUTES les fonctionnalités et cas limites de l'exigence.
        8. Assure-toi d'ÉPUISER tous les scénarios possibles, ne laisse aucun cas de test manquant."""

    else:
        # Construire le prompt en anglais
        if format_type == "default":
            instruction = f"""
            Generate test cases for the following requirement using the default format as shown in the example below.
            Each test case should include functional and non-functional scenarios, with preconditions, 
            numbered steps,  expected results, and gives the maximum number of scenarios that can cover the entire requirement.
            
            {"Functional context: " + context if context else ""}
            
            Here is an example of the expected format:
            
            {exemple_format}
            
            Now, generate detailed test cases in this same format for the following requirement:
            
            {requirements}
            """
        elif format_type == "gherkin":
            instruction = f"""
            Generate test cases for the following requirement using the Gherkin format (Given, When, Then).
            Each test case should include functional and non-functional scenarios  and gives the maximum number of scenarios that can cover the entire requirement.
            
            {"Functional context: " + context if context else ""}
            
            {"Example of expected format: " + example_case if example_case else ""}
            
            Requirement: {requirements}
            """
        else:  # format_type == "custom"
            instruction = f"""
            Generate test cases for the following requirement using exactly the custom format provided as an example.
            Strictly respect the structure and style of the provided example  and gives the maximum number of scenarios that can cover the entire requirement.
            
            {"Functional context: " + context if context else ""}
            
            Here is an example of the expected format:
            
            {exemple_format}
            
            Now, generate detailed test cases in this same custom format for the following requirement:
            
            {requirements}
            """
        
        system_prompt = """You are an expert software tester who generates high-quality test cases.
        Your job is to analyze requirements and produce comprehensive, exhaustive, and accurate test scenarios.
        For each requirement, cover all functional and non-functional aspects (performance, security, accessibility, etc.).

        STRICT FORMATTING RULES:
        1. For each SCENARIO, clearly indicate whether it is a "FUNCTIONAL" or "NON-FUNCTIONAL" case.
        2. Use appropriate line breaks for better readability.
        3. STEPS must be numbered and indented with a tab, one step per line.
        4. PRECONDITIONS must be clearly separated from the rest.
        5. EXPECTED RESULTS must be clearly highlighted, using a line break after "Expected result:".
        6. Make sure to bold (**) the section titles (Scenario, Precondition, Steps, Expected result).
        7. Generate the MAXIMUM number of scenarios possible to cover ALL the features and edge cases of the requirement.
        8. Make sure to EXHAUST all possible scenarios, don't leave any test cases missing."""

    
    # Appel direct à l'API Claude via requests
    import requests
    import time
    
    # Logique de retry avec backoff exponentiel
    max_retries = 3
    retry_delay = 2  # secondes
    
    for attempt in range(max_retries):
        try:
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers=headers,
                json={
                    "model": "claude-3-5-haiku-20241022",
                    "max_tokens": 8000,
                    "temperature": 0.2,
                    "system": system_prompt,
                    "messages": [
                        {
                            "role": "user", 
                            "content": instruction
                        }
                    ]
                }
            )
            
            if response.status_code == 200:
                try:
                    content = response.json().get("content", [])
                    text_content = ""
                    for item in content:
                        if item.get("type") == "text":
                            text_content += item.get("text", "")
                    return text_content if text_content else "Pas de texte dans la réponse"
                except Exception as e:
                    return f"Erreur lors du traitement de la réponse: {str(e)}"
            # Si API surchargée, attendre et réessayer
            elif response.status_code == 529:
                if attempt < max_retries - 1:
                    wait_time = retry_delay * (2 ** attempt)
                    time.sleep(wait_time)
                    continue
                else:
                    return f"L'API est actuellement surchargée. Veuillez réessayer plus tard."
            else:
                return f"Erreur de l'API: {response.status_code} - {response.text}"
        
        except Exception as e:
            return f"Erreur lors de l'appel à l'API Claude: {str(e)}"
# Fonction pour le chatbot qui corrige/modifie les résultats
def chat_with_results(user_query, test_cases, conversation_history):
    import os
    from dotenv import load_dotenv
    import langdetect
    
    # Charger les variables d'environnement
    load_dotenv()
    
    # Récupérer la clé API depuis les variables d'environnement
    api_key = get_api_key()
    
    if not api_key:
        return "Erreur: Clé API non trouvée dans les variables d'environnement. Vérifiez votre fichier .env"
    
    # Configuration pour l'API
    headers = {
        "anthropic-api-key": api_key,
        "content-type": "application/json",
        "anthropic-version": "2023-06-01",
        "x-api-key": api_key
    }
    
    # Détection de la langue
    try:
        detected_lang = langdetect.detect(test_cases)
        # Par défaut, on considère le français
        lang = "fr"
        
        # Si l'anglais est détecté, on adapte les instructions
        if detected_lang == "en":
            lang = "en"
    except:
        # En cas d'erreur de détection, on reste en français par défaut
        lang = "fr"
    
    # Construire l'historique des messages pour l'API
    messages = []
    
    # Ajouter l'historique de conversation existant
    for entry in conversation_history:
        messages.append({
            "role": entry["role"],
            "content": entry["content"]
        })
    
    # Ajouter le message actuel de l'utilisateur
    messages.append({
        "role": "user",
        "content": user_query
    })
    
    # Personnaliser le prompt système selon la langue
    if lang == "fr":
        system_prompt = f"""Tu es un assistant spécialisé dans la correction et l'amélioration des cas de test.
        Ton travail consiste à aider l'utilisateur à améliorer, corriger ou modifier les cas de test qui ont été générés.
        
        Voici les cas de test actuels sur lesquels tu dois travailler :
        
        {test_cases}
        
        RÈGLES ABSOLUES À SUIVRE:
        1. CONSERVE TOUJOURS TOUS LES SCÉNARIOS ORIGINAUX dans ta réponse, même quand tu ajoutes ou modifies un seul scénario.
        2. JAMAIS fournir uniquement les scénarios ajoutés ou modifiés - tu dois toujours fournir le document COMPLET.
        3. Quand l'utilisateur demande de "changer le scénario X", tu modifies ce scénario tout en conservant tous les autres scénarios.
        4. Quand l'utilisateur demande d'"ajouter des scénarios", tu ajoutes les nouveaux scénarios APRÈS tous les scénarios existants en vérfiant la numération de scenario.
        5. Si tu modifies ET ajoutes des scénarios, assure-toi que TOUS les scénarios (originaux modifiés + nouveaux) apparaissent dans ta réponse.
        6. La numérotation des scénarios doit rester cohérente (Scenario 1, Scenario 2, etc.).
        7. Assure-toi de préserver la bonne structure avec des retours à la ligne appropriés pour une meilleure lisibilité.
        
        FORMAT DE RÉPONSE REQUIS:
        - Au début de ta réponse, ajoute exactement: MODIFIED:START
        - Ensuite, inclus TOUS les cas de test (originaux + modifiés + nouveaux)
        - À la fin des cas de test, ajoute exactement: MODIFIED:END
        - Après ces balises, explique brièvement les modifications effectuées"""
    else:
        system_prompt = f"""You are an assistant specialized in correcting and improving test cases.
        Your job is to help the user improve, correct, or modify the test cases that were generated.
        
        Here are the current test cases you need to work on:
        
        {test_cases}
        
        ABSOLUTE RULES TO FOLLOW:
        1. ALWAYS KEEP ALL ORIGINAL SCENARIOS in your response, even when you add or modify a single scenario.
        2. NEVER provide only the added or modified scenarios - you must always provide the COMPLETE document.
        3. When the user asks to "change scenario X", you modify that scenario while keeping all other scenarios.
        4. When the user asks to "add scenarios", you add the new scenarios AFTER all existing scenarios.
        5. If you modify AND add scenarios, ensure that ALL scenarios (original modified + new) appear in your response.
        6. The numbering of scenarios must remain consistent (Scenario 1, Scenario 2, etc.).
        7. Make sure to preserve the proper structure with appropriate line breaks for better readability.
        
        REQUIRED RESPONSE FORMAT:
        - At the beginning of your response, add exactly: MODIFIED:START
        - Then, include ALL test cases (original + modified + new)
        - At the end of the test cases, add exactly: MODIFIED:END
        - After these tags, briefly explain the modifications made"""
    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json={
                "model": "claude-3-5-haiku-20241022",
                "max_tokens": 8000,
                "temperature": 0.3,
                "system": system_prompt,
                "messages": messages
            }
        )
        
        if response.status_code == 200:
            try:
                content = response.json().get("content", [])
                text_content = ""
                for item in content:
                    if item.get("type") == "text":
                        text_content += item.get("text", "")
                
                # Extraction des cas de test modifiés si présents
                if "MODIFIED:START" in text_content and "MODIFIED:END" in text_content:
                    start_index = text_content.find("MODIFIED:START") + len("MODIFIED:START")
                    end_index = text_content.find("MODIFIED:END")
                    if start_index > 0 and end_index > start_index:
                        st.session_state.modified_test_cases = text_content[start_index:end_index].strip()
                
                # Nettoyer le texte pour l'affichage
                display_text = text_content.replace("MODIFIED:START", "").replace("MODIFIED:END", "")
                return display_text if display_text else "Pas de texte dans la réponse"
            except Exception as e:
                return f"Erreur lors du traitement de la réponse: {str(e)}"
        else:
            return f"Erreur de l'API: {response.status_code} - {response.text}"
        
    except Exception as e:
        return f"Erreur lors de l'appel à l'API Claude: {str(e)}"
# Titre de l'application
st.title("Assistant IA - Générateur des Cas de test")
st.markdown("---")

# Initialiser les variables de session
if 'test_cases' not in st.session_state:
    st.session_state.test_cases = ""
    
if 'requirements' not in st.session_state:
    st.session_state.requirements = ""
    
if 'format_type' not in st.session_state:
    st.session_state.format_type = "default"
    
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
    
if 'modified_test_cases' not in st.session_state:
    st.session_state.modified_test_cases = ""

if 'chat_count' not in st.session_state:
    st.session_state.chat_count = 0

if 'context_input' not in st.session_state:
    st.session_state.context_input = ""

if 'edit_mode' not in st.session_state:
    st.session_state.edit_mode = False

if 'edited_content' not in st.session_state:
    st.session_state.edited_content = ""


# Cette fonction sera appelée quand l'utilisateur clique sur le bouton d'édition
def toggle_edit_mode():
    if st.session_state.edit_mode:
        # Si on sort du mode édition, sauvegarder le contenu édité
        st.session_state.modified_test_cases = st.session_state.edited_content
    else:
        # Si on entre en mode édition, initialiser le contenu éditable
        if st.session_state.modified_test_cases:
            st.session_state.edited_content = st.session_state.modified_test_cases
        else:
            st.session_state.edited_content = st.session_state.test_cases
    
    # Basculer le mode d'édition
    st.session_state.edit_mode = not st.session_state.edit_mode


# NOUVELLE STRUCTURE DE L'APPLICATION - 2 colonnes
col_input, col_output = st.columns([1, 1], gap="large")

# COLONNE 1 - Entrées utilisateur
with col_input:
    st.subheader("📝 Paramètres")
    
    # Section Contexte - Obligatoire
    with st.expander("Contexte Fonctionnel", expanded=True):
        context_tab1, context_tab2 = st.tabs(["Saisie manuelle", "Upload de fichier"])
        
        with context_tab1:
            context_input = st.text_area("Description générale de l'application (*)", 
                               value=st.session_state.context_input,
                               height=100)
            
            # Sauvegarde du contexte dans l'état de session
            if context_input:
                st.session_state.context_input = context_input
        
        with context_tab2:
            context_file = st.file_uploader("Choisir un fichier décrivant le contexte de l’application (*)", type=["txt", "docx", "pdf"], key="context_file")
            if context_file is not None:
                try:
                    if context_file.name.endswith('.docx'):
                        st.session_state.context_input = extract_text_from_docx(context_file)
                    elif context_file.name.endswith('.pdf'):
                        context_file.seek(0)
                        st.session_state.context_input = extract_text_from_pdf(context_file)
                    else:
                        context_file.seek(0)
                        st.session_state.context_input = context_file.getvalue().decode("utf-8")
                    
                    st.success(f"Contexte chargé depuis '{context_file.name}'")
                except Exception as e:
                    st.error(f"Erreur lors du chargement: {str(e)}")
    
    # Section Exigence
    with st.expander("Exigences/ User Stories", expanded=True):
        requirements_input = st.text_area("Saisir l’exigence/ la US", height=150)
        if requirements_input:
            st.session_state.requirements = requirements_input
    
    # Section Format
    with st.expander("Format des Cas de Test", expanded=True):
        format_option = st.radio(
            "Choisir le format :",
            ("Par Défaut", "Gherkin (Given When Then)", "Personnalisé (Basé sur l’exemple)"),
            captions=["Format standard des cas de test", "Format Gherkin", "Format basé sur l'exemple fourni"]
        )
        
        if format_option == "Par Défaut":
            st.session_state.format_type = "default"
        elif format_option == "Gherkin (Given When Then)":
            st.session_state.format_type = "gherkin"
        else:
            st.session_state.format_type = "custom"
        
        example_case = st.text_area(
            "Fournir le 1er Cas de test au format attendu" + ("(Obligatoire)" if format_option == "Personnalisé (Basé sur l’exemple)" else " (Optionnel)"),
            height=150
        )
    
    # Bouton de génération
    generate_btn = st.button("Générer les Cas de Test", use_container_width=True)
    
    # Style CSS personnalisé pour cette colonne
    st.markdown("""
    <style>
        div[data-testid="stExpander"] {
            background-color: #f8f9fa;
            border-radius: 10px;
            margin-bottom: 15px;
            padding: 5px;
        }
    </style>
    """, unsafe_allow_html=True)
    
# COLONNE 2 - Résultats et interactions
# Ajouter au début du script après les imports
if 'edit_mode' not in st.session_state:
    st.session_state.edit_mode = False

if 'edited_content' not in st.session_state:
    st.session_state.edited_content = ""

# Cette fonction sera appelée quand l'utilisateur clique sur le bouton d'édition
def toggle_edit_mode():
    if st.session_state.edit_mode:
        # Si on sort du mode édition, sauvegarder le contenu édité
        st.session_state.modified_test_cases = st.session_state.edited_content
    else:
        # Si on entre en mode édition, initialiser le contenu éditable
        if st.session_state.modified_test_cases:
            st.session_state.edited_content = st.session_state.modified_test_cases
        else:
            st.session_state.edited_content = st.session_state.test_cases
    
    # Basculer le mode d'édition
    st.session_state.edit_mode = not st.session_state.edit_mode

# Remplacer la section d'affichage des résultats dans col_output
with col_output:
    st.subheader("📋 Résultat")
    
    # Bouton d'édition avec icône de crayon
    edit_btn_col, download_info_col = st.columns([1, 3])
    
    with edit_btn_col:
        edit_btn = st.button("✏️ Éditer", on_click=toggle_edit_mode, 
                            use_container_width=True,
                            help="Cliquez pour éditer les résultats")
    
    with download_info_col:
        if st.session_state.modified_test_cases or st.session_state.test_cases:
            st.info("Les modifications seront incluses dans les fichiers téléchargés.")
    
    # Zone d'affichage ou d'édition des résultats
    if st.session_state.edit_mode:
        # Mode édition avec textarea
        edited_content = st.text_area(
            "Éditez les résultats ci-dessous:",
            value=st.session_state.edited_content,
            height=400,
            key="result_editor"
        )
        st.session_state.edited_content = edited_content
        
        # Bouton pour sauvegarder les modifications
        save_col, cancel_col = st.columns(2)
        with save_col:
            save_btn = st.button("💾 Sauvegarder les modifications", 
                               use_container_width=True,
                               on_click=toggle_edit_mode)
        with cancel_col:
            if st.button("❌ Annuler", use_container_width=True):
                st.session_state.edit_mode = False
                st.rerun()
    else:
        # Mode affichage normal
        result_container = st.container(height=400, border=True)
        
        # Afficher le contenu modifié en priorité s'il existe, sinon afficher les résultats originaux
        with result_container:
            if st.session_state.modified_test_cases:
                st.markdown(st.session_state.modified_test_cases)
            elif st.session_state.test_cases:
                st.markdown(st.session_state.test_cases)
            else:
                st.info("Les cas de test générés apparaîtront ici.")
    
    # Boutons d'exportation
    if st.session_state.modified_test_cases or st.session_state.test_cases or st.session_state.edited_content:
        col_pdf, col_docx = st.columns(2)
        
        with col_pdf:
            # Utiliser le contenu modifié s'il existe, sinon utiliser les résultats originaux
            if st.session_state.edit_mode:
                content_to_export = st.session_state.edited_content
            else:
                content_to_export = st.session_state.modified_test_cases if st.session_state.modified_test_cases else st.session_state.test_cases
            
            pdf_data = create_pdf(content_to_export)
            pdf_html = create_download_link(pdf_data, "pdf", "cas_de_test.pdf")
            st.markdown(pdf_html, unsafe_allow_html=True)
        
        with col_docx:
            if st.session_state.edit_mode:
                content_to_export = st.session_state.edited_content
            else:
                content_to_export = st.session_state.modified_test_cases if st.session_state.modified_test_cases else st.session_state.test_cases
            
            docx_data = create_docx(content_to_export)
            docx_html = create_download_link(docx_data, "docx", "cas_de_test.docx")
            st.markdown(docx_html, unsafe_allow_html=True)

    # Chatbot pour les modifications (limite augmentée à 10 échanges)
    if st.session_state.test_cases:
        st.markdown("---")
        st.subheader("💬 Assistant d’Amélioration (limité à 10 échanges)")
        
        # Afficher le compteur d'échanges restants
        remaining_exchanges = 10 - st.session_state.chat_count
        st.info(f"Échanges restants: {remaining_exchanges}/10")
        
        if st.session_state.chat_count >= 10:
            st.warning("Vous avez atteint la limite de 10 échanges. Vous ne pouvez plus faire de modifications.")
            
            # Afficher l'historique des messages
            with st.expander("Voir l'historique des conversations", expanded=False):
                for message in st.session_state.chat_history:
                    if message["role"] == "user":
                        st.markdown(f"**Vous:** {message['content']}")
                        st.markdown("---")
                    else:
                        st.markdown(f"**Assistant:** {message['content'].replace('MODIFIED:START', '').replace('MODIFIED:END', '')}")
                        st.markdown("---")
        else:
            # Zone de saisie du chat
            user_message = st.text_area("Demander des modifications ou des améliorations:", 
                                         placeholder="Exemple: Ajouter un scénario de test pour...", 
                                         height=100,
                                         key="user_message")
            
            # Bouton pour envoyer le message
            chat_btn = st.button("Envoyer la demande", use_container_width=True)
            
            # Traitement de l'envoi du message
            if chat_btn and user_message:
                # Compteur d'échanges
                st.session_state.chat_count += 1
                
                # Ajouter le message de l'utilisateur à l'historique
                st.session_state.chat_history.append({
                    "role": "user",
                    "content": user_message
                })
                
                # Obtenir la réponse du chatbot
                content_to_modify = st.session_state.modified_test_cases if st.session_state.modified_test_cases else st.session_state.test_cases
                with st.spinner("L'assistant réfléchit..."):
                    response = chat_with_results(
                        user_message, 
                        content_to_modify,
                        st.session_state.chat_history
                    )
                    
                    # Ajouter la réponse à l'historique (sans les marqueurs)
                    clean_response = response.replace('MODIFIED:START', '').replace('MODIFIED:END', '')
                    st.session_state.chat_history.append({
                        "role": "assistant",
                        "content": clean_response
                    })
                
                # Recharger la page pour afficher les résultats mis à jour
                st.rerun()
            
            
# Traitement du bouton de génération
if generate_btn:
    # Vérifier si le contexte est fourni (nouvelle condition obligatoire)
    if not st.session_state.context_input:
        st.error("Le contexte fonctionnel est obligatoire. Veuillez saisir un contexte ou uploader un fichier.")
    elif not st.session_state.requirements:
        st.error("Veuillez saisir une exigence.")
    elif st.session_state.format_type == "custom" and not example_case.strip():
        st.error("Pour utiliser le format personnalisé, vous devez fournir un exemple de cas de test.")
    else:
        with st.spinner("Génération des cas de test en cours..."):
            try:
                # Réinitialiser l'état lors d'une nouvelle génération
                st.session_state.modified_test_cases = ""
                st.session_state.chat_history = []
                st.session_state.chat_count = 0
                
                # Générer les cas de test
                st.session_state.test_cases = generate_test_cases(
                    st.session_state.requirements,
                    st.session_state.format_type,
                    st.session_state.context_input,
                    example_case
                )
                
                # Recharger la page pour afficher les résultats
                st.rerun()
            except Exception as e:
                st.error(f"Erreur lors de la génération: {str(e)}")

# Ajouter du CSS personnalisé
st.markdown("""
<style>
    .stButton button {
        background-color: #0066cc;
        color: white;
        font-weight: bold;
    }
    .stButton button:hover {
        background-color: #004d99;
    }
    .css-1aumxhk {
        background-color: #f8f9fa;
    }
    .stTextArea label {
        font-weight: bold;
    }
    .stRadio label {
        font-weight: bold;
    }
    
    /* Styles améliorés pour l'interface en 2 colonnes */
    [data-testid="column"] {
        background-color: #ffffff;
        padding: 10px;
        border-radius: 10px;
    }
    
    /* Style pour le conteneur de résultats */
    [data-testid="stVerticalBlock"] > div[style] > div[data-testid="stVerticalBlock"] {
        background-color: #ffffff;
        border-radius: 5px;
        padding: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)
